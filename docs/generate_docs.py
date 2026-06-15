# -*- coding: utf-8 -*-
"""
Generador de la Documentación Técnica de DesposteApp (.docx).
Construye un documento Word profesional en español.

Versión 2 del documento: refleja la arquitectura vigente del sistema
(costeo ABC, consulta SIPSA en vivo sin persistencia, cinco tablas,
restablecimiento de contraseñas por el administrador).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ----------------------------------------------------------------------------
# Paleta corporativa
# ----------------------------------------------------------------------------
BURGUNDY   = RGBColor(0x6B, 0x1E, 0x1E)   # color principal
DARK       = RGBColor(0x22, 0x22, 0x22)   # texto fuerte
GRAY       = RGBColor(0x55, 0x55, 0x55)   # texto secundario
ACCENT     = RGBColor(0xB5, 0x65, 0x1D)   # acento ámbar
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "6B1E1E"
SUBHEAD_FILL = "8C3B3B"
ZEBRA_FILL  = "F4ECEC"
CODE_FILL   = "F2F2F2"
NOTE_FILL   = "FBF3E6"

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

doc = Document()

# ----------------------------------------------------------------------------
# Helpers de bajo nivel
# ----------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tcPr.append(m)


def _no_space(paragraph, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def add_field(paragraph, field_code):
    """Inserta un campo de Word (TOC, PAGE, etc.)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = field_code
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = ""
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fldChar1); r.append(instr); r.append(fldChar2); r.append(t); r.append(fldChar3)
    return run


# ----------------------------------------------------------------------------
# Estilos base
# ----------------------------------------------------------------------------

def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1s = document.styles["Heading 1"]
    h1s.font.name = BODY_FONT
    h1s.font.size = Pt(18)
    h1s.font.bold = True
    h1s.font.color.rgb = BURGUNDY
    h1s.paragraph_format.space_before = Pt(18)
    h1s.paragraph_format.space_after = Pt(8)
    h1s.paragraph_format.keep_with_next = True

    h2s = document.styles["Heading 2"]
    h2s.font.name = BODY_FONT
    h2s.font.size = Pt(14)
    h2s.font.bold = True
    h2s.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    h2s.paragraph_format.space_before = Pt(12)
    h2s.paragraph_format.space_after = Pt(4)
    h2s.paragraph_format.keep_with_next = True

    h3s = document.styles["Heading 3"]
    h3s.font.name = BODY_FONT
    h3s.font.size = Pt(12)
    h3s.font.bold = True
    h3s.font.color.rgb = ACCENT
    h3s.paragraph_format.space_before = Pt(8)
    h3s.paragraph_format.space_after = Pt(2)
    h3s.paragraph_format.keep_with_next = True


def body(text="", bold=False, italic=False, color=None, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(parts, space_after=6, align=None):
    """parts: lista de (texto, {bold,italic,mono,color})"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for text, fmt in parts:
        r = p.add_run(text)
        fmt = fmt or {}
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        if fmt.get("mono"):
            r.font.name = MONO_FONT
            r.font.size = Pt(9.5)
        if fmt.get("color"):
            r.font.color.rgb = fmt["color"]
    return p


def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = DARK
    r2 = p.add_run(text)
    return p


def numbered_list(items):
    """Lista numerada manual (reinicia en 1 en cada llamada) con sangría francesa."""
    for i, text in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        rn = p.add_run(f"{i}.  ")
        rn.bold = True
        rn.font.color.rgb = BURGUNDY
        p.add_run(text)


def code_block(lines):
    """Bloque monoespaciado sombreado dentro de una tabla de 1 celda."""
    if isinstance(lines, str):
        lines = lines.split("\n")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, CODE_FILL)
    _set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    cell.paragraphs[0].text = ""
    first = True
    for ln in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        _no_space(p)
        r = p.add_run(ln if ln else "")
        r.font.name = MONO_FONT
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    _set_table_borders(tbl, color="DDDDDD")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def note_box(title, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, NOTE_FILL)
    _set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    p = cell.paragraphs[0]
    _no_space(p, after=2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = ACCENT
    p2 = cell.add_paragraph()
    _no_space(p2)
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    _set_table_borders(tbl, color="E5C99A")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _set_table_borders(table, color="BBBBBB", size="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def data_table(headers, rows, widths=None, header_fill=HEADER_FILL, zebra=True, font_size=9.5):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # cabecera
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_bg(hdr[i], header_fill)
        _set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        _no_space(p)
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(font_size)
    # filas
    for ridx, row in enumerate(rows):
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            _set_cell_margins(cells[i])
            if zebra and ridx % 2 == 1:
                _set_cell_bg(cells[i], ZEBRA_FILL)
            p = cells[i].paragraphs[0]
            _no_space(p)
            # soporte para (texto, mono)
            mono = False
            text = val
            if isinstance(val, tuple):
                text, mono = val
            r = p.add_run(str(text))
            r.font.size = Pt(font_size)
            if mono:
                r.font.name = MONO_FONT
                r.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)
    _set_table_borders(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def page_break():
    doc.add_page_break()


# ----------------------------------------------------------------------------
# Pie de página con numeración
# ----------------------------------------------------------------------------

def add_footer(document):
    section = document.sections[-1]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DesposteApp · Documentación Técnica · ")
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY
    add_field(p, "PAGE")
    r2 = p.add_run(" de ")
    r2.font.size = Pt(8); r2.font.color.rgb = GRAY
    add_field(p, "NUMPAGES")
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY


configure_styles(doc)
for s in doc.sections:
    s.top_margin = Cm(2.3)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.4)
    s.right_margin = Cm(2.4)

# ============================================================================
# PORTADA
# ============================================================================

def cover():
    doc.add_paragraph().paragraph_format.space_after = Pt(40)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIVERSIDAD SANTIAGO DE CALI")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GRAY
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Facultad de Ingeniería · Ingeniería de Sistemas")
    r2.font.size = Pt(11); r2.font.color.rgb = GRAY
    p2.paragraph_format.space_after = Pt(60)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run("DesposteApp")
    rt.bold = True; rt.font.size = Pt(40); rt.font.color.rgb = BURGUNDY
    t.paragraph_format.space_after = Pt(4)

    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rst = st.add_run("Sistema de Optimización de Precios para el Desposte de Ganado Bovino")
    rst.font.size = Pt(15); rst.font.color.rgb = DARK
    st.paragraph_format.space_after = Pt(2)

    st2 = doc.add_paragraph(); st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rst2 = st2.add_run("Documentación Técnica del Sistema")
    rst2.italic = True; rst2.font.size = Pt(13); rst2.font.color.rgb = ACCENT
    st2.paragraph_format.space_after = Pt(70)

    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ficha = [
        ("Proyecto", "DesposteApp — Prototipo de software"),
        ("Versión del documento", "2.0"),
        ("Versión del sistema", "2.0.0"),
        ("Fecha", datetime.date.today().strftime("%d/%m/%Y")),
        ("Estado", "Prototipo funcional (entorno académico)"),
        ("Autores", "Juan Esteban Montilla Rayo · Juan Pablo Rodríguez Becerra · Rafael Ángel Davalos Villegas"),
        ("Audiencia", "Cliente técnico / equipo de desarrollo"),
        ("Confidencialidad", "Documento de uso interno"),
    ]
    for k, v in ficha:
        cells = tbl.add_row().cells
        _set_cell_margins(cells[0]); _set_cell_margins(cells[1])
        _set_cell_bg(cells[0], "EFE7E7")
        pk = cells[0].paragraphs[0]; _no_space(pk)
        rk = pk.add_run(k); rk.bold = True; rk.font.size = Pt(10); rk.font.color.rgb = BURGUNDY
        pv = cells[1].paragraphs[0]; _no_space(pv)
        rv = pv.add_run(v); rv.font.size = Pt(10)
        cells[0].width = Inches(1.9); cells[1].width = Inches(4.4)
    _set_table_borders(tbl, color="D8C9C9")
    page_break()

cover()

# ============================================================================
# TABLA DE CONTENIDO
# ============================================================================
h1("Tabla de Contenido")
body("Para actualizar la numeración de páginas y las entradas, haga clic derecho sobre la tabla y seleccione «Actualizar campos» (o presione F9) en Microsoft Word.",
     italic=True, color=GRAY, size=9.5)
toc_p = doc.add_paragraph()
add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u')
page_break()

# ============================================================================
# 1. INTRODUCCIÓN
# ============================================================================
h1("1. Introducción")

h2("1.1 Propósito del documento")
body("Este documento describe el diseño, la arquitectura, los componentes y los requisitos de "
     "DesposteApp en su versión actual. Está pensado como referencia para quien deba evaluar, "
     "mantener o extender el sistema: explica cómo está construido, cómo se calculan los precios, "
     "cómo se integra con la fuente de datos del DANE y cómo se instala y opera. También dedica una "
     "sección a dejar por escrito lo que el sistema no hace, porque conocer los límites de un "
     "prototipo es tan importante como conocer sus funciones.")
body("Esta es la segunda versión del documento. Reemplaza a la anterior e incorpora los cambios "
     "más relevantes del sistema desde entonces: el modelo de costeo por actividades (ABC) para "
     "distribuir los costos entre cortes, la simplificación del módulo SIPSA (los datos de mercado "
     "se consultan en vivo y ya no se almacenan), la reducción del modelo de datos a cinco tablas y "
     "el restablecimiento de contraseñas por parte del administrador.")

h2("1.2 Alcance del sistema")
body("DesposteApp es un prototipo de software que apoya la toma de decisiones de precios en el "
     "proceso de desposte de ganado bovino. El sistema calcula precios de venta sugeridos por corte "
     "a partir de tres insumos: (1) el costo de compra del animal y los costos del proceso, "
     "distribuidos entre los cortes mediante costeo ABC; (2) los precios de referencia del mercado "
     "mayorista publicados por el DANE a través del servicio SIPSA; y (3) un margen de ganancia "
     "objetivo configurable. Además gestiona el registro de animales, cortes y costos, ofrece un "
     "panel de indicadores y administra usuarios con control de acceso por roles.")

h2("1.3 Audiencia")
bullet("Responsable de evaluar técnicamente la solución antes de su adopción.", bold_prefix="Cliente técnico: ")
bullet("Personas encargadas de mantener, corregir o ampliar el sistema.", bold_prefix="Equipo de desarrollo: ")
bullet("Personal que despliega y administra el entorno donde corre la aplicación.", bold_prefix="Operación: ")

h2("1.4 Glosario y definiciones")
data_table(
    ["Término", "Definición"],
    [
        ["Desposte", "Proceso de despiece de la canal del animal en cortes comerciales individuales."],
        ["Canal", "Cuerpo del animal sacrificado, sin vísceras, piel ni extremidades, listo para despiece."],
        ["Peso vivo", "Peso del animal en pie, antes del sacrificio."],
        ["Peso canal", "Peso del animal una vez faenado (la canal)."],
        ["Rendimiento en canal", "Relación porcentual entre el peso canal y el peso vivo."],
        ["Merma", "Pérdida de peso entre la canal y la suma de cortes vendibles (huesos, recortes, etc.)."],
        ["Corte", "Pieza comercial obtenida del despiece (p. ej. lomo fino, cadera, costilla)."],
        ["Costeo ABC", "Costeo basado en actividades: asigna los costos a los productos según los recursos que cada uno consume."],
        ["Inductor de costo", "Criterio con el que un costo se reparte entre los cortes (peso, horas-hombre, kWh, etc.)."],
        ["Factor de complejidad", "Ponderador que refleja cuánto trabajo de proceso exige un corte respecto al promedio."],
        ["SIPSA", "Sistema de Información de Precios y Abastecimiento del Sector Agropecuario (DANE)."],
        ["DANE", "Departamento Administrativo Nacional de Estadística de Colombia."],
        ["Margen de ganancia", "Porcentaje de utilidad objetivo sobre el precio de venta."],
        ["JWT", "JSON Web Token; estándar para tokens de autenticación firmados."],
        ["ORM", "Object-Relational Mapping; mapeo objeto-relacional (SQLAlchemy)."],
        ["SOAP", "Protocolo de intercambio de mensajes basado en XML usado por el servicio del DANE."],
        ["SPA", "Single Page Application; aplicación web de una sola página."],
    ],
    widths=[1.9, 4.4],
)
page_break()

# ============================================================================
# 2. DESCRIPCIÓN GENERAL
# ============================================================================
h1("2. Descripción General del Sistema")

h2("2.1 Contexto y problema")
body("En las plantas de desposte y comercializadoras de carne bovina, la fijación de precios suele "
     "realizarse de forma manual y empírica. Esto dificulta incorporar de manera consistente el costo "
     "real del proceso, el rendimiento en canal de cada animal y las variaciones del mercado mayorista. "
     "Un problema adicional, menos evidente, es que no todos los cortes cuestan lo mismo de producir: "
     "un lomo fino exige más trabajo de limpieza y porcionado que la carne molida, pero un reparto de "
     "costos por simple peso ignora esa diferencia. DesposteApp aborda ambos problemas: sistematiza el "
     "cálculo integrando costos, rendimiento y precios de referencia oficiales, y distribuye los costos "
     "entre cortes según la intensidad de proceso de cada uno.")

h2("2.2 Objetivos del sistema")
bullet("Centralizar el registro de animales, cortes y costos del proceso de desposte.")
bullet("Calcular el costo real por corte mediante costeo basado en actividades (ABC).")
bullet("Incorporar precios de referencia del mercado (DANE-SIPSA) al cálculo de precios.")
bullet("Sugerir precios de venta por corte aplicando un margen objetivo configurable.")
bullet("Ofrecer un panel de indicadores para apoyar la toma de decisiones.")
bullet("Controlar el acceso mediante autenticación y roles de usuario.")

h2("2.3 Características principales")
data_table(
    ["Módulo", "Descripción"],
    [
        ["Autenticación", "Inicio de sesión con JWT en cookie HttpOnly, gestión de perfil y cambio de contraseña."],
        ["Animales", "Registro y gestión de animales con peso vivo, peso canal, calidad y precio de compra."],
        ["Cortes", "Gestión de cortes por animal, con categoría, peso, factor de complejidad y precios calculados."],
        ["Costos", "Registro de costos por animal o globales, cada uno con su inductor de costo ABC."],
        ["SIPSA", "Consulta en vivo de precios de referencia del mercado bovino del DANE, con promedios y tendencia por corte."],
        ["Análisis", "Panel de indicadores, costo por kilogramo y cálculo de precios sugeridos por animal."],
        ["Usuarios", "Administración de usuarios, roles y restablecimiento de contraseñas (solo perfil ADMIN)."],
    ],
    widths=[1.6, 4.7],
)

h2("2.4 Roles de usuario")
data_table(
    ["Rol", "Permisos"],
    [
        ["ADMIN", "Acceso completo a todos los módulos, incluida la administración de usuarios."],
        ["ESTANDAR", "Acceso a operación (animales, cortes, costos, SIPSA, análisis); sin gestión de usuarios."],
    ],
    widths=[1.5, 4.8],
)
note_box("Usuarios iniciales",
         "Al arrancar por primera vez, el backend crea automáticamente los usuarios definidos en las "
         "variables de entorno (un administrador y un operador). Las credenciales no están escritas "
         "en el código fuente: se leen de la configuración del entorno, y si no se define la "
         "contraseña de un usuario inicial, ese usuario no se crea.")

h2("2.5 Flujo de trabajo de extremo a extremo")
body("El uso típico del sistema sigue una secuencia que va desde la compra del animal hasta el precio "
     "sugerido de cada corte. El recorrido completo, tal como lo vive el usuario, es el siguiente:")
numbered_list([
    "Inicio de sesión: el usuario se autentica en /login; el backend emite la cookie de sesión.",
    "Registro del animal (módulo Animales): se captura el código, tipo, peso vivo, peso canal, "
    "calidad y precio de compra. El sistema calcula el rendimiento en canal si dispone de ambos pesos.",
    "Registro de cortes (módulo Cortes): se registran los cortes obtenidos del despiece con su peso. "
    "El sistema les asigna automáticamente un factor de complejidad según el nombre del corte.",
    "Registro de costos (módulo Costos): se ingresan los costos del proceso (transporte, sacrificio, "
    "energía, mano de obra...), asociados al animal o al fondo global, indicando el inductor con el "
    "que cada costo debe repartirse entre los cortes.",
    "Consulta de mercado (módulo SIPSA): la aplicación consulta en vivo los precios mayoristas del "
    "DANE y muestra promedios, rangos y tendencia por corte. Esta consulta es informativa; el cálculo "
    "de precios la realiza de nuevo por su cuenta.",
    "Cálculo de precios (módulo Análisis): el usuario selecciona el animal y el margen objetivo. El "
    "backend distribuye los costos entre los cortes (costeo ABC), mezcla el resultado con la "
    "referencia de mercado y devuelve el precio sugerido, el mínimo viable y el margen real por corte.",
    "Persistencia de resultados: cada cálculo actualiza los campos del corte y deja un registro en la "
    "tabla de precios, de modo que queda trazabilidad de cómo se determinó cada precio.",
    "Seguimiento (Dashboard): el panel muestra los totales, el costo promedio por kilogramo, el margen "
    "promedio y la tabla de referencia SIPSA.",
])
page_break()

# ============================================================================
# 3. ARQUITECTURA
# ============================================================================
h1("3. Arquitectura del Sistema")

h2("3.1 Visión general")
body("DesposteApp adopta una arquitectura cliente-servidor de tres capas, organizada como un "
     "monorepo con dos aplicaciones independientes: un frontend SPA (Single Page Application) en React "
     "y un backend de API REST en FastAPI. La persistencia se realiza sobre una base de datos "
     "PostgreSQL (en la nube, mediante Supabase). El backend consume además un servicio externo SOAP "
     "del DANE para obtener precios de referencia del mercado; estos datos se procesan en memoria en "
     "cada consulta y no se almacenan en la base de datos.")

code_block([
    "┌──────────────────────────────────────────────────────────────────────┐",
    "│                          NAVEGADOR (Cliente)                          │",
    "│   React 18 + Vite + React Router + Axios  (SPA, puerto 5173)          │",
    "└───────────────────────────────┬──────────────────────────────────────┘",
    "                                 │  HTTP /api/v1/*  (cookie HttpOnly JWT)",
    "                                 │  proxy de Vite en desarrollo",
    "┌───────────────────────────────▼──────────────────────────────────────┐",
    "│                       BACKEND — API REST (FastAPI)                     │",
    "│   Rutas → Dependencias (Auth/Roles) → Servicios → Modelos ORM         │",
    "│   JWT (HS256) · bcrypt · SlowAPI (rate limit) · CORS  (puerto 8001)   │",
    "└──────────────┬───────────────────────────────┬───────────────────────┘",
    "               │ SQLAlchemy                     │ zeep (SOAP, en vivo)",
    "┌──────────────▼─────────────┐      ┌───────────▼──────────────────────┐",
    "│   PostgreSQL (Supabase)    │      │   DANE — SIPSA  (Web Service)     │",
    "│   5 tablas relacionales    │      │   con fallback a datos demo       │",
    "└────────────────────────────┘      └───────────────────────────────────┘",
])

h2("3.2 Estructura monorepo")
body("El repositorio agrupa las dos aplicaciones bajo una misma raíz:")
code_block([
    "tesis/",
    "├── backend/            API REST en FastAPI (Python)",
    "│   ├── main.py         Punto de entrada y configuración de la aplicación",
    "│   ├── requirements.txt",
    "│   └── app/",
    "│       ├── routes/       Routers FastAPI (capa de presentación)",
    "│       ├── services/     Lógica de negocio (auth, análisis de precios)",
    "│       ├── models/       Modelos ORM SQLAlchemy (tablas)",
    "│       ├── schemas/      Esquemas Pydantic (validación E/S)",
    "│       ├── dependencies/ Dependencias de autenticación y autorización",
    "│       ├── database/     Motor, sesión y Base declarativa",
    "│       ├── sipsa/        Cliente SOAP DANE y procesamiento de datos",
    "│       └── limiter.py    Limitador de peticiones (SlowAPI)",
    "└── frontend/           SPA en React + Vite (JavaScript)",
    "    ├── vite.config.js   Configuración y proxy /api",
    "    └── src/",
    "        ├── pages/       Una página por ruta",
    "        ├── components/  Navbar, Toast, rutas protegidas, tarjetas KPI",
    "        ├── context/     AuthContext (estado de sesión)",
    "        ├── services/    api.js (instancia Axios)",
    "        └── App.jsx      Definición de rutas",
])

h2("3.3 Arquitectura por capas del backend")
body("El backend separa responsabilidades en capas bien definidas:")
data_table(
    ["Capa", "Responsabilidad", "Ubicación"],
    [
        ["Presentación (Routers)", "Define los endpoints REST, valida entradas y delega en servicios.", ("app/routes/", True)],
        ["Validación (Schemas)", "Esquemas Pydantic para validar y serializar peticiones y respuestas.", ("app/schemas/", True)],
        ["Dependencias", "Autenticación, extracción de token y control de roles.", ("app/dependencies/", True)],
        ["Negocio (Services)", "Reglas de cálculo de precios, costeo ABC y autenticación.", ("app/services/", True)],
        ["Datos (Models)", "Modelos ORM SQLAlchemy mapeados a tablas.", ("app/models/", True)],
        ["Integración", "Cliente SOAP del DANE y procesamiento de datos de mercado.", ("app/sipsa/", True)],
    ],
    widths=[1.5, 3.6, 1.4],
)

h2("3.4 Comunicación frontend-backend")
bullet("En desarrollo, Vite expone un proxy que redirige toda petición a /api hacia el backend "
       "(http://localhost:8001), evitando problemas de CORS.", bold_prefix="Proxy: ")
bullet("Axios se configura con baseURL /api/v1 y withCredentials: true, de modo que la cookie de "
       "sesión viaja automáticamente en cada petición.", bold_prefix="Cliente HTTP: ")
bullet("El token JWT se almacena en una cookie HttpOnly emitida por el backend; el JavaScript del "
       "navegador no necesita (ni puede) leerlo, lo que reduce el riesgo de robo por XSS.", bold_prefix="Sesión: ")
bullet("Ante una respuesta 401, el interceptor de Axios limpia el estado local y redirige a /login.",
       bold_prefix="Manejo de sesión expirada: ")

h2("3.5 Ciclo de vida de una petición")
numbered_list([
    "El usuario interactúa con una página React, que invoca un método del servicio Axios (api.js).",
    "La petición sale hacia /api/v1/... con la cookie de sesión; el proxy de Vite la reenvía al backend.",
    "FastAPI resuelve el router correspondiente; la dependencia de autenticación valida el JWT y el rol.",
    "El router invoca la capa de servicio, que ejecuta la lógica de negocio usando los modelos ORM.",
    "SQLAlchemy traduce las operaciones a SQL contra PostgreSQL y devuelve los resultados.",
    "La respuesta se serializa con un esquema Pydantic y regresa al frontend en formato JSON.",
])
page_break()

# ============================================================================
# 4. STACK Y REQUISITOS
# ============================================================================
h1("4. Stack Tecnológico y Requisitos")

h2("4.1 Backend")
data_table(
    ["Tecnología", "Versión de referencia", "Función"],
    [
        ["Python", "3.11+", "Lenguaje del backend."],
        ["FastAPI", "0.111", "Framework web/API REST."],
        ["Uvicorn", "0.30", "Servidor ASGI de ejecución."],
        ["SQLAlchemy", "2.0", "ORM y acceso a base de datos."],
        ["Pydantic", "2.7", "Validación y serialización de datos."],
        ["psycopg2-binary", "2.9", "Driver de PostgreSQL."],
        ["python-jose", "3.3", "Firma y verificación de JWT (HS256)."],
        ["bcrypt", "4.1", "Hash de contraseñas."],
        ["zeep", "4.2", "Cliente SOAP para el servicio del DANE."],
        ["pandas / numpy", "2.2 / 1.26", "Procesamiento y limpieza de datos de mercado."],
        ["slowapi", "0.1.9", "Limitación de tasa de peticiones (rate limiting)."],
        ["python-dotenv", "1.0", "Carga de variables de entorno desde .env."],
    ],
    widths=[1.8, 1.7, 2.8],
)
note_box("Nota sobre versiones",
         "Las versiones indicadas corresponden a las fijadas en requirements.txt. En entornos con "
         "Python muy reciente (p. ej. 3.13/3.14) puede ser necesario emplear versiones más actuales de "
         "numpy y pandas, que disponen de paquetes precompilados (wheels) compatibles.")

h2("4.2 Frontend")
data_table(
    ["Tecnología", "Versión de referencia", "Función"],
    [
        ["Node.js", "18+", "Entorno de ejecución para construir el frontend."],
        ["React", "18.3", "Librería de interfaz de usuario."],
        ["Vite", "5.2", "Servidor de desarrollo y empaquetador."],
        ["React Router", "6.23", "Enrutamiento de la SPA."],
        ["Axios", "1.7", "Cliente HTTP hacia la API."],
        ["lucide-react", "1.17", "Iconografía SVG de la interfaz."],
    ],
    widths=[1.8, 1.7, 2.8],
)

h2("4.3 Base de datos")
body("Motor relacional PostgreSQL. El prototipo está configurado para usar Supabase (PostgreSQL "
     "gestionado en la nube) a través de su «Session Pooler». El backend crea automáticamente el "
     "esquema (las cinco tablas) en el arranque y aplica de forma idempotente las columnas del modelo "
     "ABC sobre instalaciones previas, sin necesidad de migraciones manuales.")

h2("4.4 Servicios externos")
body("El sistema consume el servicio web SIPSA del DANE (SOAP) para obtener precios de referencia del "
     "mercado bovino. Cuando el servicio no responde, el sistema recurre automáticamente a un conjunto "
     "de datos de demostración con precios realistas, de modo que la aplicación sigue siendo funcional. "
     "Los datos de mercado se consultan y procesan en cada petición; no se persisten localmente.")

h2("4.5 Requisitos de software y hardware")
data_table(
    ["Requisito", "Mínimo recomendado"],
    [
        ["Sistema operativo", "Linux, macOS o Windows 10/11"],
        ["Python", "3.11 o superior"],
        ["Node.js", "18 o superior"],
        ["Memoria RAM", "4 GB (8 GB recomendado para desarrollo)"],
        ["Conectividad", "Acceso a internet (Supabase y servicio DANE-SIPSA)"],
        ["Navegador", "Chrome, Edge o Firefox en versión reciente"],
    ],
    widths=[2.2, 4.1],
)
page_break()

# ============================================================================
# 5. MODELO DE DATOS
# ============================================================================
h1("5. Modelo de Datos")

h2("5.1 Diagrama entidad-relación")
body("El modelo se compone de cinco entidades. Un animal agrupa varios cortes y varios costos; cada "
     "corte mantiene un histórico de precios calculados. Los usuarios son una entidad independiente "
     "dedicada a la autenticación. Los datos de mercado SIPSA no se almacenan: se consultan en vivo "
     "al servicio del DANE cuando se necesitan.")
code_block([
    "                         ┌───────────────┐",
    "                         │   usuarios    │   (autenticación / roles)",
    "                         └───────────────┘",
    "",
    "   ┌───────────┐ 1     N ┌───────────┐ 1     N ┌───────────┐",
    "   │  animales │─────────│  cortes   │─────────│  precios  │",
    "   └───────────┘         └───────────┘         └───────────┘",
    "        │ 1",
    "        │",
    "        │ N",
    "   ┌───────────┐",
    "   │  costos   │   (animal_id nulo = costo del fondo global)",
    "   └───────────┘",
])
body("Relaciones definidas con SQLAlchemy:", bold=True, space_after=2)
bullet("animales → cortes: uno a muchos, con borrado en cascada.")
bullet("animales → costos: uno a muchos, con borrado en cascada (animal_id puede ser nulo para costos globales).")
bullet("cortes → precios: uno a muchos, con borrado en cascada (histórico de cálculos).")

h2("5.2 Detalle de tablas")

h3("Tabla: animales")
data_table(
    ["Campo", "Tipo", "Restricciones / Descripción"],
    [
        [("id", True), "Integer", "PK, autoincremental."],
        [("codigo", True), "String(50)", "Único, obligatorio. Identificador del animal."],
        [("tipo", True), "String(20)", "Obligatorio. Por defecto «BOVINO»."],
        [("raza", True), "String(100)", "Opcional."],
        [("peso_vivo", True), "Float", "Obligatorio, > 0. Peso en pie."],
        [("peso_canal", True), "Float", "Opcional al registrar; necesario para calcular precios."],
        [("rendimiento_canal", True), "Float", "Calculado: (peso_canal / peso_vivo) × 100. Nulo si falta el peso canal."],
        [("calidad", True), "String(20)", "Por defecto «PRIMERA»."],
        [("precio_compra", True), "Float", "Obligatorio, > 0. Costo de adquisición."],
        [("fecha_sacrificio", True), "DateTime", "Opcional."],
        [("fecha_registro", True), "DateTime", "Automática al crear."],
        [("notas", True), "String(500)", "Opcional."],
    ],
    widths=[1.8, 1.3, 3.5],
)

h3("Tabla: cortes")
data_table(
    ["Campo", "Tipo", "Restricciones / Descripción"],
    [
        [("id", True), "Integer", "PK."],
        [("animal_id", True), "Integer", "FK → animales.id. Obligatorio."],
        [("nombre", True), "String(100)", "Obligatorio. Nombre del corte."],
        [("categoria", True), "String(50)", "PREMIUM / ESTANDAR / ECONOMICO. Por defecto «ESTANDAR»."],
        [("peso_kg", True), "Float", "Obligatorio. Peso del corte."],
        [("factor_complejidad", True), "Float", "Intensidad de proceso del corte (ponderador ABC). Se asigna automáticamente según el nombre; 1,0 = corte promedio."],
        [("costo_unitario", True), "Float", "Costo ABC por kg, escrito en el último cálculo de precios."],
        [("precio_sugerido", True), "Float", "Resultado del cálculo de precios."],
        [("precio_mercado_sipsa", True), "Float", "Precio de referencia SIPSA aplicado."],
        [("margen_ganancia", True), "Float", "Margen real resultante."],
        [("activo", True), "Boolean", "Por defecto verdadero."],
        [("fecha_registro", True), "DateTime", "Automática al crear."],
    ],
    widths=[1.9, 1.2, 3.5],
)

h3("Tabla: costos")
data_table(
    ["Campo", "Tipo", "Restricciones / Descripción"],
    [
        [("id", True), "Integer", "PK."],
        [("animal_id", True), "Integer", "FK → animales.id. Nulo = costo global."],
        [("concepto", True), "String(200)", "Obligatorio. Ej.: transporte, sacrificio."],
        [("categoria", True), "String(100)", "Opcional."],
        [("valor", True), "Float", "Obligatorio. Monto del costo en pesos."],
        [("unidad", True), "String(50)", "Por defecto «por_animal»."],
        [("inductor", True), "String(30)", "Inductor ABC: KG, HORAS_HOMBRE, KWH, M3_REFRIG o FIJO. Por defecto «KG»."],
        [("fecha_registro", True), "DateTime", "Automática al crear."],
        [("notas", True), "String(500)", "Opcional."],
    ],
    widths=[1.8, 1.3, 3.5],
)

h3("Tabla: precios (histórico de cálculos)")
data_table(
    ["Campo", "Tipo", "Restricciones / Descripción"],
    [
        [("id", True), "Integer", "PK."],
        [("corte_id", True), "Integer", "FK → cortes.id. Obligatorio."],
        [("precio_costo_unitario", True), "Float", "Costo ABC por kg empleado en el cálculo."],
        [("precio_sugerido", True), "Float", "Precio de venta sugerido."],
        [("margen_objetivo", True), "Float", "Margen objetivo solicitado (por defecto 25 %)."],
        [("precio_sipsa_referencia", True), "Float", "Precio de mercado usado como referencia."],
        [("precio_minimo_viable", True), "Float", "Precio mínimo para alcanzar el margen."],
        [("precio_maximo_mercado", True), "Float", "Precio máximo de mercado observado."],
        [("fecha_calculo", True), "DateTime", "Automática al calcular."],
        [("activo", True), "Boolean", "Por defecto verdadero."],
    ],
    widths=[2.0, 1.1, 3.5],
)

h3("Tabla: usuarios")
data_table(
    ["Campo", "Tipo", "Restricciones / Descripción"],
    [
        [("id", True), "Integer", "PK."],
        [("nombre", True), "String(100)", "Obligatorio."],
        [("email", True), "String(150)", "Único, obligatorio."],
        [("username", True), "String(50)", "Único, obligatorio."],
        [("hashed_password", True), "String(255)", "Contraseña con hash bcrypt. Nunca en texto plano."],
        [("rol", True), "Enum", "ADMIN o ESTANDAR."],
        [("activo", True), "Boolean", "Habilitación de la cuenta."],
        [("password_version", True), "Integer", "Versiona la contraseña para invalidar tokens previos."],
        [("fecha_creacion", True), "DateTime", "Automática al crear."],
        [("ultimo_acceso", True), "DateTime", "Marca del último inicio de sesión."],
    ],
    widths=[1.9, 1.1, 3.6],
)
page_break()

# ============================================================================
# 6. LÓGICA DE NEGOCIO
# ============================================================================
h1("6. Lógica de Negocio: Costeo ABC y Cálculo de Precios")
body("El corazón funcional del sistema reside en el servicio de análisis "
     "(app/services/analisis_service.py) y en el procesador de datos de mercado "
     "(app/sipsa/processor.py). El cálculo combina dos ideas: primero, repartir el costo total del "
     "animal entre sus cortes de manera proporcional a lo que cada corte realmente consume (costeo "
     "ABC); segundo, contrastar ese costo con el precio del mercado mayorista para llegar a un precio "
     "sugerido razonable.")

h2("6.1 Rendimiento en canal")
rich([("Si se dispone del peso canal, el rendimiento se calcula como ", {}),
      ("rendimiento = (peso_canal / peso_vivo) × 100", {"mono": True}),
      (". Cuando el peso canal no se ha registrado, el rendimiento queda nulo: es un dato real que se "
       "mide tras el sacrificio, y el sistema prefiere no inventar un valor por defecto.", {})])

h2("6.2 Costo directo por kilogramo de canal")
body("Este indicador, visible en el panel, agrega el precio de compra y los costos asignados al "
     "animal, dividido por el peso de la canal:", space_after=4)
code_block([
    "costo_adicional = Σ (valor de cada costo asignado al animal)",
    "",
    "costo_por_kg = (precio_compra + costo_adicional) / peso_canal",
])
body("Si el animal no tiene peso canal registrado, el indicador queda nulo. Conviene precisar que "
     "este es un promedio informativo del animal completo; el costo que se usa para fijar precios es "
     "el costo ABC por corte, que se describe a continuación.", italic=True, color=GRAY, size=9.5)

h2("6.3 Costeo basado en actividades (ABC)")
body("Repartir todos los costos por peso trataría igual un kilogramo de lomo fino y uno de carne "
     "molida, cuando el primero exige bastante más trabajo de proceso. Para corregirlo, cada costo "
     "registrado declara un inductor que determina cómo se distribuye entre los cortes del animal:")
data_table(
    ["Inductor", "Se reparte por", "Ejemplo de costo"],
    [
        [("KG", True), "Peso del corte", "Transporte"],
        [("FIJO", True), "Peso del corte", "Tarifa de sacrificio"],
        [("HORAS_HOMBRE", True), "Peso × factor de complejidad", "Mano de obra de desposte"],
        [("KWH", True), "Peso × factor de complejidad", "Energía de refrigeración"],
        [("M3_REFRIG", True), "Peso × factor de complejidad", "Espacio en cuartos fríos"],
    ],
    widths=[1.5, 2.2, 2.6],
)
body("El factor de complejidad refleja cuánto trabajo exige cada corte respecto al promedio y se "
     "asigna automáticamente a partir del nombre, para que el usuario no tenga que estimar un número:")
data_table(
    ["Corte", "Factor", "Corte", "Factor"],
    [
        ["Lomo fino", "1,5", "Bola negra / Lagarto", "1,1"],
        ["Punta de anca", "1,3", "Costilla / Brazo", "1,0"],
        ["Lomo de aguja / Cadera / Muchacho", "1,2", "Pecho", "0,9"],
        ["Corte no reconocido", "1,0", "Molida", "0,8"],
    ],
    widths=[2.6, 0.8, 2.1, 0.8],
)
body("Con esos elementos, la distribución se realiza así:", space_after=4)
code_block([
    "peso_total = Σ peso de los cortes del animal",
    "base_total = Σ (peso × factor_complejidad)",
    "",
    "# Participan los costos del animal y los del fondo global (animal_id nulo)",
    "costo_por_peso    = Σ costos con inductor KG o FIJO",
    "costo_por_proceso = Σ costos con inductor HORAS_HOMBRE, KWH o M3_REFRIG",
    "",
    "para cada corte:",
    "    parte_peso    = peso / peso_total",
    "    parte_proceso = (peso × factor_complejidad) / base_total",
    "",
    "    costo_material    = precio_compra × parte_peso",
    "    costo_actividades = costo_por_peso × parte_peso",
    "                      + costo_por_proceso × parte_proceso",
    "",
    "    costo_unitario = (costo_material + costo_actividades) / peso",
])
body("Un detalle importante: el precio de compra se reparte solo entre los cortes vendibles "
     "registrados. Como la suma de cortes pesa menos que la canal (huesos, recortes, merma), el costo "
     "de esa diferencia queda absorbido por los cortes aprovechables, que es lo que ocurre en la "
     "práctica comercial. Por la misma razón, el costo ABC por kilogramo resulta algo mayor que el "
     "costo directo por kilogramo de canal de la sección 6.2.")

h2("6.4 Generación del precio sugerido")
body("Para cada corte, el costo unitario ABC se transforma en precio sugerido en cuatro pasos:")
code_block([
    "# 1. Saneamiento del margen (si margen ≥ 100 %, se normaliza a 30 %)",
    "precio_base = costo_unitario / (1 - margen_objetivo / 100)",
    "",
    "# 2. Ajuste por categoría del corte",
    "factor = { PREMIUM: 1.15,  ESTANDAR: 1.00,  ECONOMICO: 0.88 }",
    "precio_ajustado = precio_base × factor[categoria]",
    "",
    "# 3. Mezcla con el precio de mercado SIPSA (si está disponible)",
    "si hay precio_sipsa:",
    "    precio_final = precio_ajustado × 0.6 + precio_sipsa × 0.4",
    "si no:",
    "    precio_final = precio_ajustado",
    "",
    "# 4. Margen real resultante",
    "margen_real = (precio_final - costo_unitario) / precio_final × 100",
])
body("Interpretación de cada paso:", bold=True, space_after=2)
bullet("Es el precio mínimo que recupera el costo y alcanza el margen objetivo; se reporta como "
       "«precio mínimo viable».", bold_prefix="Precio base: ")
bullet("Premia los cortes premium (+15 %) y abarata los económicos (−12 %), reflejando su valor "
       "comercial relativo.", bold_prefix="Factor de categoría: ")
bullet("Pondera 60 % el precio basado en costos y 40 % el precio de mercado, para que la sugerencia "
       "no se aleje de la realidad del mercado mayorista. Si no hay referencia para ese corte, el "
       "precio queda determinado solo por costos.", bold_prefix="Mezcla con SIPSA: ")
bullet("Mide el margen efectivo del precio final; si la mezcla con el mercado lo dejó por debajo del "
       "objetivo, la interfaz lo resalta para que el usuario lo revise.", bold_prefix="Margen real: ")
body("El cálculo exige que el animal tenga peso canal registrado; en caso contrario la API responde "
     "con un error 400 indicándolo. Cada ejecución actualiza los campos del corte (costo unitario, "
     "precio sugerido, referencia de mercado y margen) e inserta un registro en la tabla precios, "
     "conservando la trazabilidad de cada cálculo.")

h2("6.5 Ejemplo numérico")
body("Animal comprado en $2.500.000, con peso vivo de 450 kg y peso canal de 250 kg (rendimiento "
     "55,6 %). Costos registrados: transporte $150.000 (KG), sacrificio $120.000 (FIJO), energía de "
     "frío $80.000 (KWH) y mano de obra $200.000 (HORAS_HOMBRE). Para simplificar, el despiece se "
     "reduce a tres cortes que suman 220 kg vendibles:")
data_table(
    ["Corte", "Peso", "Factor", "Costo material", "Costo actividades", "Costo unitario"],
    [
        ["Lomo fino (PREMIUM)", "10 kg", "1,5", "$113.636", "$32.563", "$14.620/kg"],
        ["Cadera (ESTANDAR)", "60 kg", "1,2", "$681.818", "$171.027", "$14.214/kg"],
        ["Molida (ECONOMICO)", "150 kg", "0,8", "$1.704.545", "$346.410", "$13.673/kg"],
    ],
    widths=[1.9, 0.7, 0.7, 1.2, 1.3, 1.1],
    font_size=9,
)
body("La suma de costos distribuidos ($3.050.000) coincide con el total real (compra más costos), y "
     "el lomo fino termina con el costo unitario más alto pese a ser el corte más liviano, porque "
     "concentra más trabajo de proceso. Con ese costo, el precio de la cadera (margen objetivo del "
     "25 % y referencia SIPSA de $22.000/kg) se obtiene así:")
data_table(
    ["Paso", "Operación", "Resultado"],
    [
        ["Precio base (mínimo viable)", ("14214 / (1 - 0.25)", True), "$18.952"],
        ["Ajuste categoría (×1,00)", ("18952 × 1.00", True), "$18.952"],
        ["Mezcla SIPSA", ("18952 × 0.6 + 22000 × 0.4", True), "$20.171"],
        ["Margen real", ("(20171 - 14214) / 20171 × 100", True), "≈ 29,5 %"],
    ],
    widths=[2.1, 2.6, 1.6],
)
page_break()

# ============================================================================
# 7. INTEGRACIÓN SIPSA
# ============================================================================
h1("7. Integración con DANE-SIPSA")

h2("7.1 Qué es SIPSA")
body("SIPSA es el Sistema de Información de Precios y Abastecimiento del Sector Agropecuario, "
     "administrado por el DANE. Publica precios mayoristas de productos agropecuarios —entre ellos la "
     "carne bovina— en distintas centrales de abasto del país. DesposteApp consume este servicio para "
     "anclar sus precios sugeridos a referencias oficiales de mercado.")

h2("7.2 Cliente SOAP")
rich([("El cliente (", {}), ("app/sipsa/client.py", {"mono": True}),
      (") emplea la librería ", {}), ("zeep", {"mono": True}),
      (" para invocar el servicio web SOAP del DANE. Filtra los productos bovinos mediante un conjunto "
       "de palabras clave (res, bovina, lomo, costilla, cadera, etc.) y normaliza cada registro a una "
       "estructura común: producto, ciudad, mercado, precios (promedio, mínimo, máximo), unidad y fechas.", {})])

h2("7.3 Estrategia de tolerancia a fallos (4 niveles)")
body("La obtención de datos está diseñada para no fallar de cara al usuario. Se intentan, en orden, "
     "distintos métodos del servicio y, como último recurso, un conjunto de datos de demostración:")
data_table(
    ["Nivel", "Fuente", "Método / contenido"],
    [
        ["1", "DANE — por ciudad", ("promediosSipsaCiudad()", True)],
        ["2", "DANE — parcial", ("promediosSipsaParcial()  (min/máx por ubicación)", True)],
        ["3", "DANE — semanal mayorista", ("promediosSipsaSemanaMadr()", True)],
        ["4", "Datos de demostración", "12 cortes × 4 mercados × 12 semanas, precios realistas Colombia"],
    ],
    widths=[0.8, 2.0, 3.5],
)
note_box("Consulta en vivo, sin persistencia",
         "Los datos de mercado se obtienen y procesan en memoria cada vez que se necesitan (consulta "
         "del módulo SIPSA, panel o cálculo de precios) y no se guardan en la base de datos. Esto "
         "simplifica el sistema y evita acumular duplicados, a cambio de no construir un histórico "
         "propio de precios; esa decisión se discute en la sección de limitaciones.")

h2("7.4 Procesamiento y limpieza de datos")
rich([("El procesador (", {}), ("app/sipsa/processor.py", {"mono": True}),
      (") aplica una secuencia de limpieza estadística sobre los datos recibidos:", {})])
bullet("Conversión de precios a numérico y descarte de valores nulos o no positivos.", bold_prefix="Saneamiento: ")
bullet("Eliminación de valores atípicos mediante el rango intercuartílico (regla 1,5 × IQR).", bold_prefix="Outliers: ")
bullet("Agrupación por corte (mediante un mapa de palabras clave con normalización de tildes y "
       "mayúsculas) y cálculo de promedio, mínimo, máximo y desviación estándar.", bold_prefix="Agregación: ")

h2("7.5 Cálculo de la tendencia")
body("La columna de tendencia que se muestra en el panel y en el módulo SIPSA (SUBIENDO, BAJANDO o "
     "ESTABLE) se calcula por corte de la siguiente manera: se ordenan las observaciones por fecha, "
     "se toman las cuatro más recientes y se compara el precio de la primera con el de la última. Si "
     "la variación porcentual supera el +3 % la tendencia es SUBIENDO; si cae por debajo del −3 %, "
     "BAJANDO; en la franja intermedia, ESTABLE. Cuando hay menos de tres observaciones con fecha, el "
     "sistema devuelve ESTABLE por defecto. Es deliberadamente un método simple —compara extremos, no "
     "ajusta una regresión—, suficiente como señal visual rápida pero no como análisis estadístico de "
     "series de tiempo.")
page_break()

# ============================================================================
# 8. API REST
# ============================================================================
h1("8. Referencia de la API REST")
body("Todos los endpoints se exponen bajo el prefijo /api/v1. Salvo el inicio de sesión, todas las "
     "rutas requieren autenticación. La documentación interactiva (Swagger UI) está disponible en "
     "/docs únicamente cuando la variable DEBUG está activada.")

h2("8.1 Autenticación y usuarios — /api/v1/auth")
data_table(
    ["Método", "Ruta", "Descripción", "Acceso"],
    [
        [("POST", True), ("/auth/login", True), "Inicia sesión y emite la cookie JWT.", "Público (5/min)"],
        [("POST", True), ("/auth/logout", True), "Cierra sesión y elimina la cookie.", "Autenticado"],
        [("GET", True), ("/auth/me", True), "Devuelve los datos del usuario actual.", "Autenticado"],
        [("POST", True), ("/auth/cambiar-password", True), "Cambia la contraseña propia (invalida tokens previos).", "Autenticado"],
        [("GET", True), ("/auth/usuarios", True), "Lista los usuarios.", "ADMIN"],
        [("POST", True), ("/auth/usuarios", True), "Crea un usuario.", "ADMIN"],
        [("PUT", True), ("/auth/usuarios/{id}", True), "Actualiza nombre, correo, rol o estado; opcionalmente restablece la contraseña (mínimo 8 caracteres, invalida las sesiones del usuario).", "ADMIN"],
        [("DELETE", True), ("/auth/usuarios/{id}", True), "Elimina un usuario.", "ADMIN"],
    ],
    widths=[0.8, 1.9, 3.0, 1.0],
    font_size=9,
)

h2("8.2 Animales — /api/v1/animales")
data_table(
    ["Método", "Ruta", "Descripción"],
    [
        [("POST", True), ("/animales/", True), "Crea un animal (calcula el rendimiento en canal si hay datos)."],
        [("GET", True), ("/animales/", True), "Lista animales (paginado: skip, limit)."],
        [("GET", True), ("/animales/{id}", True), "Obtiene un animal por su identificador."],
        [("PUT", True), ("/animales/{id}", True), "Actualiza un animal y recalcula su rendimiento."],
        [("DELETE", True), ("/animales/{id}", True), "Elimina un animal (y sus cortes/costos en cascada)."],
    ],
    widths=[0.9, 2.2, 3.7],
)

h2("8.3 Cortes — /api/v1/cortes")
data_table(
    ["Método", "Ruta", "Descripción"],
    [
        [("POST", True), ("/cortes/", True), "Crea un corte (asigna automáticamente el factor de complejidad)."],
        [("GET", True), ("/cortes/", True), "Lista cortes (paginado)."],
        [("GET", True), ("/cortes/animal/{id}", True), "Lista los cortes de un animal."],
        [("GET", True), ("/cortes/{id}", True), "Obtiene un corte."],
        [("PUT", True), ("/cortes/{id}", True), "Actualiza un corte."],
        [("DELETE", True), ("/cortes/{id}", True), "Elimina un corte."],
        [("GET", True), ("/cortes/{id}/precios", True), "Histórico de precios calculados del corte."],
    ],
    widths=[0.9, 2.3, 3.6],
)

h2("8.4 Costos — /api/v1/costos")
data_table(
    ["Método", "Ruta", "Descripción"],
    [
        [("POST", True), ("/costos/", True), "Crea un costo (por animal o global) con su inductor ABC."],
        [("GET", True), ("/costos/", True), "Lista costos (paginado)."],
        [("GET", True), ("/costos/animal/{id}", True), "Lista los costos de un animal."],
        [("GET", True), ("/costos/{id}", True), "Obtiene un costo."],
        [("PUT", True), ("/costos/{id}", True), "Actualiza un costo."],
        [("DELETE", True), ("/costos/{id}", True), "Elimina un costo."],
    ],
    widths=[0.9, 2.3, 3.6],
)

h2("8.5 SIPSA — /api/v1/sipsa")
data_table(
    ["Método", "Ruta", "Descripción"],
    [
        [("GET", True), ("/sipsa/consultar", True), "Consulta en vivo los precios de mercado (parámetros de fecha opcionales; devuelve hasta 200 registros)."],
        [("GET", True), ("/sipsa/promedios", True), "Promedios, rangos, desviación y tendencia por corte, calculados sobre la consulta en vivo."],
    ],
    widths=[0.9, 2.2, 3.7],
)

h2("8.6 Análisis — /api/v1/analisis")
data_table(
    ["Método", "Ruta", "Descripción"],
    [
        [("GET", True), ("/analisis/dashboard", True), "Indicadores globales (totales, costo y margen promedio)."],
        [("POST", True), ("/analisis/calcular-precios/{id}", True), "Calcula y persiste los precios de los cortes de un animal (parámetro margen; exige peso canal)."],
        [("GET", True), ("/analisis/costo-kg/{id}", True), "Devuelve el costo directo por kilogramo de un animal."],
    ],
    widths=[0.9, 2.7, 3.2],
)
page_break()

# ============================================================================
# 9. SEGURIDAD
# ============================================================================
h1("9. Seguridad")

h2("9.1 Autenticación basada en JWT")
body("El inicio de sesión válido genera un JSON Web Token firmado con el algoritmo HS256. El token "
     "incluye el usuario (sub), el rol, el identificador y la versión de contraseña (pwd_v), con una "
     "expiración configurable mediante la variable TOKEN_EXPIRE_MINUTES (60 minutos por defecto).")

h2("9.2 Cookie HttpOnly")
body("El token se entrega al navegador en una cookie HttpOnly con SameSite=Lax. Al no ser accesible "
     "desde JavaScript, se mitiga el robo de token mediante ataques XSS. En producción (DEBUG "
     "desactivado) la cookie se marca como Secure, exigiendo HTTPS. Para clientes de API y para la "
     "documentación interactiva se acepta también el token en la cabecera Authorization (Bearer).")

h2("9.3 Hash de contraseñas")
body("Las contraseñas nunca se almacenan en texto plano: se guardan con hash bcrypt y sal aleatoria. "
     "La verificación se realiza comparando el hash, sin exponer jamás la contraseña original.")

h2("9.4 Invalidación de sesiones por versión de contraseña")
body("Cada usuario tiene un campo password_version que viaja dentro del token. Cuando la contraseña "
     "cambia —sea porque el propio usuario la actualiza desde su perfil o porque un administrador la "
     "restablece desde la gestión de usuarios—, ese número se incrementa y todos los tokens emitidos "
     "antes del cambio dejan de ser aceptados de inmediato, aunque no hayan expirado.")

h2("9.5 Autorización por roles")
rich([("La dependencia ", {}), ("get_current_user", {"mono": True}),
      (" valida el token en cada petición protegida y verifica que la cuenta siga activa; la "
       "dependencia ", {}),
      ("require_admin", {"mono": True}),
      (" restringe las rutas de administración de usuarios al rol ADMIN, devolviendo 403 en caso contrario.", {})])

h2("9.6 Limitación de tasa (rate limiting)")
body("El endpoint de inicio de sesión está limitado a 5 intentos por minuto y dirección IP mediante "
     "SlowAPI, mitigando ataques de fuerza bruta.")

h2("9.7 CORS")
body("La política de CORS del backend permite únicamente los orígenes de desarrollo del frontend "
     "(localhost:5173 y localhost:3000), con credenciales habilitadas y métodos y cabeceras acotados.")

h2("9.8 Gestión de secretos")
body("La clave de firma (SECRET_KEY), la cadena de conexión a la base de datos y las credenciales de "
     "los usuarios iniciales se leen exclusivamente de variables de entorno. El archivo .env está "
     "excluido del control de versiones (.gitignore). Si falta la SECRET_KEY, el backend se niega a "
     "arrancar en lugar de usar una clave por defecto insegura.")

note_box("Recomendaciones para producción",
         "Para un despliegue productivo se recomienda: rotar la SECRET_KEY y mantenerla en un gestor de "
         "secretos; forzar HTTPS extremo a extremo; restringir CORS al dominio real; aplicar políticas "
         "de contraseña robustas; añadir registro de auditoría; y revisar las versiones de dependencias "
         "frente a vulnerabilidades conocidas.")
page_break()

# ============================================================================
# 10. FRONTEND
# ============================================================================
h1("10. Frontend (Interfaz de Usuario)")

h2("10.1 Tecnología y organización")
rich([("La interfaz es una SPA construida con React 18 y Vite. El estado de sesión se gestiona con un "
       "contexto global (", {}), ("AuthContext", {"mono": True}),
      ("), el enrutamiento con React Router 6 y la comunicación con la API mediante una instancia "
       "centralizada de Axios (", {}), ("services/api.js", {"mono": True}),
      ("). La iconografía utiliza lucide-react, una librería de iconos SVG, en lugar de caracteres "
       "decorativos, para mantener una apariencia uniforme en cualquier sistema operativo.", {})])

h2("10.2 Gestión de la sesión en el cliente")
body("Como el token vive en una cookie HttpOnly, el JavaScript del navegador no puede leerlo. El "
     "frontend guarda en localStorage únicamente una marca de sesión activa y el perfil básico del "
     "usuario (nombre, usuario y rol) para pintar la interfaz. Al recargar la página, si existe la "
     "marca de sesión, el contexto consulta /auth/me para validar que la cookie siga vigente y "
     "recuperar el perfil; si la respuesta es 401, limpia el estado y redirige al inicio de sesión.")

h2("10.3 Páginas y rutas")
data_table(
    ["Ruta", "Página", "Descripción", "Acceso"],
    [
        [("/login", True), "LoginPage", "Inicio de sesión.", "Público"],
        [("/", True), "DashboardPage", "Panel de indicadores y tabla de referencia SIPSA.", "Autenticado"],
        [("/animales", True), "AnimalesPage", "Gestión de animales.", "Autenticado"],
        [("/cortes", True), "CortesPage", "Gestión de cortes (muestra factor ABC y costos calculados).", "Autenticado"],
        [("/costos", True), "CostosPage", "Gestión de costos con inductor ABC.", "Autenticado"],
        [("/sipsa", True), "SIPSAPage", "Consulta de precios de mercado (promedios y detalle).", "Autenticado"],
        [("/analisis", True), "AnalisisPage", "Cálculo de precios y análisis por animal.", "Autenticado"],
        [("/usuarios", True), "UsuariosPage", "Administración de usuarios y restablecimiento de contraseñas.", "ADMIN"],
        [("/perfil", True), "PerfilPage", "Perfil y cambio de contraseña propia.", "Autenticado"],
    ],
    widths=[1.0, 1.4, 3.0, 0.9],
    font_size=9,
)

h2("10.4 Rutas protegidas")
rich([("El componente ", {}), ("ProtectedRoute", {"mono": True}),
      (" envuelve las vistas autenticadas y redirige a /login cuando no hay sesión. Las rutas marcadas "
       "como ", {}), ("adminOnly", {"mono": True}),
      (" (p. ej. /usuarios) exigen además el rol ADMIN. El interceptor de Axios fuerza la redirección "
       "al recibir una respuesta 401.", {})])

h2("10.5 Componentes reutilizables")
bullet("Barra de navegación principal con accesos según el rol.", bold_prefix="Navbar: ")
bullet("Control de acceso a las vistas autenticadas y de administrador.", bold_prefix="ProtectedRoute: ")
bullet("Notificaciones emergentes de éxito o error.", bold_prefix="Toast: ")
bullet("Tarjetas de indicadores para el panel.", bold_prefix="KpiCard / StatsCard: ")
page_break()

# ============================================================================
# 11. INSTALACIÓN Y DESPLIEGUE
# ============================================================================
h1("11. Instalación y Despliegue")

h2("11.1 Requisitos previos")
bullet("Python 3.11 o superior.")
bullet("Node.js 18 o superior.")
bullet("Una base de datos PostgreSQL accesible (p. ej. un proyecto en Supabase).")

h2("11.2 Configuración de la base de datos (Supabase)")
numbered_list([
    "Crear un proyecto en Supabase.",
    "En Settings → Database → Connection string, copiar la URL del «Session Pooler» (puerto 5432).",
    "Registrar dicha URL en la variable DATABASE_URL del archivo backend/.env.",
])
body("Las tablas y los usuarios iniciales se crean automáticamente al arrancar el backend.",
     italic=True, color=GRAY, size=9.5)

h2("11.3 Backend")
code_block([
    "cd backend",
    "python -m venv venv",
    "source venv/bin/activate          # Windows: venv\\Scripts\\activate",
    "pip install -r requirements.txt",
    "uvicorn main:app --reload --port 8001",
    "",
    "# Documentación interactiva (Swagger, solo con DEBUG=true):",
    "#   http://localhost:8001/docs",
])

h2("11.4 Frontend")
code_block([
    "cd frontend",
    "npm install",
    "npm run dev        # servidor de desarrollo en http://localhost:5173",
    "npm run build      # build de producción",
])
note_box("Coherencia de puertos",
         "El proxy del frontend (frontend/vite.config.js) apunta al backend en el puerto 8001. Si se "
         "ejecuta el backend en otro puerto, debe actualizarse el destino del proxy para que coincida.")

h2("11.5 Variables de entorno (backend/.env)")
data_table(
    ["Variable", "Obligatoria", "Descripción"],
    [
        [("DATABASE_URL", True), "Sí", "Cadena de conexión PostgreSQL (Supabase Session Pooler)."],
        [("SECRET_KEY", True), "Sí", "Clave de firma de los JWT. Sin ella, el backend no arranca."],
        [("DEBUG", True), "No", "true habilita /docs y cookies no seguras (desarrollo)."],
        [("TOKEN_EXPIRE_MINUTES", True), "No", "Expiración del token en minutos (por defecto 60)."],
        [("SIPSA_WSDL", True), "No", "URL del servicio SOAP del DANE."],
        [("ADMIN_USERNAME / _PASSWORD / _EMAIL", True), "No*", "Credenciales del usuario administrador inicial."],
        [("OPERADOR_USERNAME / _PASSWORD / _EMAIL", True), "No*", "Credenciales del usuario operador inicial."],
    ],
    widths=[2.6, 1.0, 2.7],
    font_size=9,
)
body("* Si no se define la contraseña de un usuario inicial, dicho usuario simplemente no se crea.",
     italic=True, color=GRAY, size=9)
page_break()

# ============================================================================
# 12. OPERACIÓN Y MANTENIMIENTO
# ============================================================================
h1("12. Operación y Mantenimiento")

h2("12.1 Arranque del sistema")
body("Al iniciar, el backend: (1) crea el esquema de base de datos si no existe; (2) aplica de forma "
     "idempotente las columnas del modelo ABC (inductor en costos; factor de complejidad y costo "
     "unitario en cortes) sobre bases de datos creadas con versiones anteriores; (3) crea los usuarios "
     "iniciales definidos por entorno; y (4) registra los manejadores de límite de tasa y CORS. El "
     "frontend se sirve por separado y se comunica con el backend a través del proxy.")

h2("12.2 Consideraciones de mantenimiento")
bullet("El esquema se crea con create_all y unos ALTER TABLE idempotentes; para cambios de modelo en "
       "producción se recomienda adoptar una herramienta de migraciones (p. ej. Alembic).",
       bold_prefix="Migraciones: ")
bullet("Se aconseja programar respaldos periódicos de la base de datos (gestionados por Supabase).",
       bold_prefix="Respaldos: ")
bullet("Conviene revisar periódicamente las versiones de las dependencias y aplicar parches de seguridad.",
       bold_prefix="Dependencias: ")
bullet("El backend registra eventos relevantes (consultas SIPSA, activación del fallback, errores) "
       "mediante el módulo de logging.", bold_prefix="Registro: ")
page_break()

# ============================================================================
# 13. LIMITACIONES
# ============================================================================
h1("13. Limitaciones y Alcance No Cubierto")
body("DesposteApp es un prototipo académico y conviene ser explícitos sobre sus fronteras. Esta "
     "sección reúne lo que el sistema no hace, las simplificaciones que asume y las decisiones de "
     "diseño cuyas consecuencias deben conocerse antes de usarlo o extenderlo. Nada de lo que sigue "
     "impide el funcionamiento del prototipo; son los puntos donde un uso en producción exigiría "
     "trabajo adicional.")

h2("13.1 Alcance funcional no cubierto")
bullet("El sistema termina en la sugerencia de precio: no gestiona inventarios, ventas, facturación, "
       "clientes ni proveedores.")
bullet("No genera reportes exportables (PDF o Excel); los resultados se consultan en pantalla.")
bullet("El análisis de precios se realiza animal por animal; no existe un cálculo por lotes ni una "
       "vista consolidada por periodo (semana, mes).")
bullet("La merma no se modela como entidad propia: la diferencia entre el peso de la canal y la suma "
       "de los cortes registrados se absorbe implícitamente en el costo de material de los cortes "
       "vendibles. Es un supuesto razonable, pero el sistema no informa cuánta merma hubo ni permite "
       "valorarla por separado (p. ej. venta de hueso o sebo).")
bullet("Maneja una sola moneda (pesos colombianos) y no contempla impuestos, descuentos ni "
       "promociones en el precio sugerido.")
bullet("No construye un histórico propio de precios de mercado: los datos SIPSA se consultan en vivo "
       "y se descartan tras usarse. Por eso no es posible graficar la evolución de un precio a lo "
       "largo de meses, y la tendencia mostrada depende de las observaciones que el servicio del DANE "
       "devuelva en esa consulta puntual.")
bullet("La interfaz está disponible solo en español y pensada para pantallas de escritorio; no se ha "
       "hecho una adaptación seria a dispositivos móviles ni una auditoría de accesibilidad.")
bullet("Está pensado para un solo negocio (una planta de desposte); no es multiempresa ni "
       "multi-sucursal.")

h2("13.2 Limitaciones del modelo de cálculo")
body("El modelo de precios funciona y es trazable, pero varios de sus parámetros son heurísticos: se "
     "fijaron por criterio propio durante el desarrollo y no fueron calibrados contra datos reales de "
     "ventas. En concreto:")
bullet("Los factores de categoría (+15 % premium, −12 % económico), la ponderación de la mezcla con "
       "el mercado (60/40) y los factores de complejidad (0,8 a 1,5) son constantes definidas en el "
       "código; no son configurables desde la interfaz ni se estimaron empíricamente.")
bullet("El costeo ABC reparte montos en pesos según el inductor declarado; no calcula tarifas a "
       "partir de cantidades medidas (horas-hombre reales, kWh del medidor). Es decir, el usuario "
       "registra «mano de obra: $200.000» y el sistema lo distribuye; no registra «8 horas a $25.000».")
bullet("La tendencia de mercado compara el primer y el último punto de las cuatro observaciones más "
       "recientes, con un umbral fijo de ±3 %. Un valor atípico en cualquiera de los dos extremos "
       "puede cambiar la etiqueta; no se ajusta una regresión ni se suaviza la serie.")
bullet("El costo directo por kilogramo del panel considera solo los costos asignados al animal, "
       "mientras que el costeo ABC incluye también los costos del fondo global. Ambos indicadores son "
       "coherentes por separado, pero no son directamente comparables entre sí.")
bullet("El cálculo exige el peso de la canal; si no se ha registrado, el sistema no estima un "
       "rendimiento teórico, simplemente no calcula. Es una decisión deliberada (no inventar datos), "
       "pero implica que el flujo se bloquea hasta tener la pesa real.")

h2("13.3 Limitaciones de los datos de mercado")
bullet("La integración depende de la disponibilidad del servicio SOAP del DANE, que en la práctica "
       "es intermitente. Cuando no responde, el sistema pasa a datos de demostración generados con "
       "una semilla fija: son verosímiles (rangos reales de Colombia 2024-2025) y estables entre "
       "ejecuciones, pero no son precios reales. La interfaz no distingue visualmente cuál de las dos "
       "fuentes está activa, de modo que el usuario debe ser consciente de esta posibilidad antes de "
       "tomar decisiones con esos valores.")
bullet("Algunos métodos del servicio solo entregan el precio promedio; en esos casos el mínimo y el "
       "máximo se estiman como ±7 % del promedio, igual que en los datos de demostración.")
bullet("El filtrado de productos bovinos se hace por palabras clave sobre el nombre del producto; "
       "puede dejar pasar productos con nombres parecidos o excluir denominaciones regionales que no "
       "estén en el mapa de cortes.")

h2("13.4 Limitaciones técnicas y de seguridad")
bullet("No hay pruebas automatizadas (unitarias ni de integración) ni un pipeline de integración "
       "continua; la verificación ha sido manual.")
bullet("El esquema se crea con create_all más unos ALTER TABLE idempotentes; no hay una herramienta "
       "de migraciones real, por lo que la evolución del modelo en una base con datos exige cuidado.")
bullet("La limitación de tasa solo cubre el inicio de sesión y se mantiene en memoria: se reinicia "
       "con el proceso y no se comparte entre réplicas.")
bullet("No hay refresh tokens (la sesión expira y hay que volver a entrar), ni recuperación de "
       "contraseña por correo, ni autenticación de dos factores, ni registro de auditoría de acciones.")
bullet("CORS está fijado a los orígenes de desarrollo (localhost); un despliegue real requiere "
       "ajustarlo, junto con HTTPS extremo a extremo.")
bullet("El acceso a la base de datos es síncrono y está dimensionado para los volúmenes bajos de un "
       "entorno académico; no se ha hecho ninguna prueba de carga.")
bullet("Solo se documenta la ejecución en modo desarrollo; no hay imagen de contenedor (Docker) ni "
       "guía de despliegue productivo.")

h2("13.5 Trabajo futuro sugerido")
body("De las limitaciones anteriores se desprende una lista natural de mejoras, ordenada por el valor "
     "que aportaría a la operación:")
numbered_list([
    "Persistir un histórico propio de precios SIPSA (con deduplicación por producto y semana) para "
    "graficar tendencias reales y mejorar el cálculo de tendencia con series de tiempo.",
    "Indicar en la interfaz si los datos de mercado provienen del DANE o del conjunto de demostración.",
    "Incorporar pruebas unitarias y de integración, y un pipeline de despliegue.",
    "Adoptar migraciones de base de datos (Alembic).",
    "Exportación de reportes (PDF/Excel) de precios y márgenes.",
    "Hacer configurables los parámetros del modelo (factores de categoría, mezcla con mercado, "
    "factores de complejidad) y, a futuro, calibrarlos con datos reales de ventas.",
    "Modelar la merma y los subproductos del desposte como parte del cálculo.",
    "Análisis consolidado por lotes y por periodo.",
])

# ----------------------------------------------------------------------------
# Cierre
# ----------------------------------------------------------------------------
doc.add_paragraph().paragraph_format.space_after = Pt(12)
closing = doc.add_paragraph(); closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = closing.add_run("— Fin del documento —")
rc.italic = True; rc.font.color.rgb = GRAY; rc.font.size = Pt(10)

add_footer(doc)

OUT = "Documentacion_Tecnica_DesposteApp.docx"
import os
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), OUT)
doc.save(out_path)
print("Documento generado:", out_path)
